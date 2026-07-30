"""
tests/test_resume_parser.py
────────────────────────────
Unit tests for services/resume_parser.py (Phase 2)

Run with:
    pytest tests/test_resume_parser.py -v
"""

from __future__ import annotations

from pathlib import Path

import pytest

from services.resume_parser import (
    ParseResult,
    ResumeParser,
    clean_text,
    extract_docx_text,
    parse_upload,
)


# ── Fixtures ───────────────────────────────────────────────────────────────────

@pytest.fixture
def parser() -> ResumeParser:
    return ResumeParser()


@pytest.fixture
def sample_docx_bytes(tmp_path: Path) -> bytes:
    """Create a minimal DOCX file in memory for testing."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_paragraph("Jane Smith")
    doc.add_paragraph("jane@example.com | San Francisco, CA")
    doc.add_paragraph("Senior Backend Engineer")
    doc.add_paragraph("Skills: Python, FastAPI, PostgreSQL, Docker")
    doc.add_paragraph("Experience: Acme Corp (2021–Present) — Backend Engineer")
    doc.add_paragraph("Education: B.S. Computer Science, Stanford University, 2018")

    out_path = tmp_path / "sample_resume.docx"
    doc.save(str(out_path))
    return out_path.read_bytes()


@pytest.fixture
def sample_docx_path(tmp_path: Path, sample_docx_bytes: bytes) -> Path:
    p = tmp_path / "sample_resume.docx"
    p.write_bytes(sample_docx_bytes)
    return p


# ── ParseResult tests ──────────────────────────────────────────────────────────

class TestParseResult:
    def test_user_facing_error_scanned(self):
        r = ParseResult(is_scanned=True)
        assert "scanned" in r.user_facing_error.lower()

    def test_user_facing_error_encrypted(self):
        r = ParseResult(is_encrypted=True)
        assert "password" in r.user_facing_error.lower()

    def test_user_facing_error_too_short(self):
        r = ParseResult(is_too_short=True, char_count=12)
        assert "12" in r.user_facing_error

    def test_user_facing_error_invalid_type(self):
        r = ParseResult(is_invalid_type=True)
        assert "unsupported" in r.user_facing_error.lower()


# ── clean_text tests ───────────────────────────────────────────────────────────

class TestCleanText:
    def test_empty_string(self):
        assert clean_text("") == ""

    def test_strips_leading_trailing_whitespace(self):
        assert clean_text("  hello  ") == "hello"

    def test_collapses_multiple_blank_lines(self):
        result = clean_text("Section A\n\n\n\n\nSection B")
        assert "\n\n\n" not in result

    def test_removes_standalone_page_numbers(self):
        text = "Work Experience\n\n2\n\nEducation"
        result = clean_text(text)
        lines = [l.strip() for l in result.splitlines() if l.strip()]
        assert "2" not in lines

    def test_removes_page_keyword(self):
        text = "Skills\n\nPage 3\n\nEducation"
        result = clean_text(text)
        assert "Page 3" not in result

    def test_preserves_bullet_points(self):
        text = "Experience\n• Developed APIs\n• Led team of 5"
        result = clean_text(text)
        assert "•" in result

    def test_preserves_section_names(self):
        text = "WORK EXPERIENCE\nSenior Engineer at Acme"
        result = clean_text(text)
        assert "WORK EXPERIENCE" in result

    def test_normalizes_unicode(self):
        result = clean_text("Pro\uFB01le")   # ﬁ ligature
        assert "fi" in result.lower() or "Profile" in result

    def test_collapses_repeated_spaces(self):
        result = clean_text("Python    FastAPI   Docker")
        assert "  " not in result


# ── parse_upload tests ─────────────────────────────────────────────────────────

class TestParseUpload:
    def test_invalid_extension_returns_failure(self):
        result = parse_upload(b"data", "resume.txt")
        assert result.success is False
        assert result.is_invalid_type is True

    def test_valid_docx_returns_success(self, sample_docx_bytes):
        result = parse_upload(sample_docx_bytes, "resume.docx")
        assert result.success is True
        assert result.char_count > 0
        assert "Jane Smith" in result.text

    def test_result_has_word_count_data(self, sample_docx_bytes):
        result = parse_upload(sample_docx_bytes, "resume.docx")
        assert result.success is True
        assert len(result.text.split()) > 5

    def test_corrupt_pdf_returns_failure(self):
        result = parse_upload(b"NOT A REAL PDF", "resume.pdf")
        assert result.success is False

    def test_too_short_pdf_returns_failure(self):
        import fitz
        doc = fitz.open()
        doc.new_page()
        pdf_bytes = doc.tobytes()
        result = parse_upload(pdf_bytes, "blank.pdf")
        assert result.success is False
        assert result.is_scanned or result.is_too_short


# ── extract_docx_text tests ────────────────────────────────────────────────────

class TestExtractDocxText:
    def test_returns_string(self, sample_docx_bytes):
        text = extract_docx_text(sample_docx_bytes)
        assert isinstance(text, str)
        assert len(text) > 0

    def test_contains_expected_content(self, sample_docx_bytes):
        text = extract_docx_text(sample_docx_bytes)
        assert "Jane Smith" in text
        assert "Python" in text

    def test_extracts_content_from_tables_and_headers(self):
        from io import BytesIO
        from docx import Document

        doc = Document()
        doc.sections[0].header.paragraphs[0].text = "Jane Smith | Portfolio"
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Skills"
        table.cell(0, 1).text = (
            "Python, FastAPI, PostgreSQL, Docker, AWS, Kubernetes, Terraform "
            "and production backend engineering experience"
        )
        buffer = BytesIO()
        doc.save(buffer)

        text = extract_docx_text(buffer.getvalue())
        assert "Jane Smith | Portfolio" in text
        assert "Python" in text
        assert "production backend engineering experience" in text


# ── ResumeParser class (backwards-compat) ─────────────────────────────────────

class TestResumeParserInit:
    def test_supported_extensions(self, parser):
        assert ".pdf" in parser.SUPPORTED_EXTENSIONS
        assert ".docx" in parser.SUPPORTED_EXTENSIONS


class TestResumeParserValidation:
    def test_raises_for_missing_file(self, parser, tmp_path):
        with pytest.raises(FileNotFoundError):
            parser.parse(tmp_path / "nonexistent.pdf")

    def test_raises_for_unsupported_extension(self, parser, tmp_path):
        txt_file = tmp_path / "resume.txt"
        txt_file.write_text("some content")
        with pytest.raises(ValueError, match="Unsupported file type"):
            parser.parse(txt_file)

    def test_raises_for_unsupported_bytes_extension(self, parser):
        with pytest.raises(ValueError, match="Unsupported file type"):
            parser.parse_bytes(b"data", ".txt")


class TestDocxParsing:
    def test_parses_docx_returns_string(self, parser, sample_docx_path):
        text = parser.parse(sample_docx_path)
        assert isinstance(text, str)
        assert len(text) > 0

    def test_docx_contains_name(self, parser, sample_docx_path):
        text = parser.parse(sample_docx_path)
        assert "Jane Smith" in text

    def test_docx_contains_skills(self, parser, sample_docx_path):
        text = parser.parse(sample_docx_path)
        assert "Python" in text

    def test_parse_bytes_docx(self, parser, sample_docx_bytes):
        text = parser.parse_bytes(sample_docx_bytes, ".docx")
        assert "Jane Smith" in text


class TestPdfParsing:
    def test_pdf_parsing_returns_text(self, parser, tmp_path):
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text(
            (72, 72),
            (
                "Jane Smith\nSenior Backend Engineer\n"
                "Python FastAPI PostgreSQL Docker AWS\n"
                "Built and operated reliable APIs used by enterprise customers.\n"
                "B.S. Computer Science"
            ),
        )
        path = tmp_path / "resume.pdf"
        path.write_bytes(doc.tobytes())
        doc.close()

        text = parser.parse(path)

        assert "Jane Smith" in text
        assert "Python" in text
