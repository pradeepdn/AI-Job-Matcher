"""
services/resume_parser.py
──────────────────────────
Phase 2 — Resume text extraction and cleaning.

This is plain Python code (not an AI agent).

Public API:
    parse_upload(file_bytes, extension)  → ParseResult
    extract_pdf_text(file_bytes)         → str
    extract_docx_text(file_bytes)        → str
    clean_text(raw_text)                 → str

Error handling:
    - Scanned / image-only PDF  → ParseResult with is_scanned=True
    - Password-protected PDF    → ParseResult with is_encrypted=True
    - Too little text extracted → ParseResult with is_too_short=True
    - Unsupported type          → raises ValueError
"""

from __future__ import annotations

import logging
import re
import unicodedata
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ── Constants ──────────────────────────────────────────────────────────────────
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}
MIN_CHARS = 100          # fewer chars than this = suspiciously short
SCANNED_THRESHOLD = 50   # avg chars per page below this → likely scanned


# ── Result dataclass ───────────────────────────────────────────────────────────

@dataclass
class ParseResult:
    """
    Outcome of parsing a resume file.

    Always check `success` before using `text`.
    """
    text: str = ""
    char_count: int = 0
    page_count: int = 0           # PDF only
    success: bool = False
    error_message: Optional[str] = None

    # Failure sub-types (mutually exclusive)
    is_scanned: bool = False
    is_encrypted: bool = False
    is_too_short: bool = False
    is_invalid_type: bool = False

    @property
    def user_facing_error(self) -> str:
        """Return a user-friendly error string for display in the UI."""
        if self.is_scanned:
            return (
                "📷 This resume appears to be a scanned image. "
                "Please upload a text-based PDF or a DOCX file. "
                "OCR support will be added in a future version."
            )
        if self.is_encrypted:
            return (
                "🔒 This PDF is password-protected. "
                "Please remove the password and re-upload."
            )
        if self.is_too_short:
            return (
                "⚠️ Very little text was extracted from this file "
                f"(only {self.char_count} characters). "
                "Make sure the file is a text-based PDF or DOCX resume."
            )
        if self.is_invalid_type:
            return (
                "❌ Unsupported file type. "
                "Please upload a PDF (.pdf) or Word document (.docx)."
            )
        return self.error_message or "❌ An unknown error occurred while parsing the file."


# ── Main entry point ───────────────────────────────────────────────────────────

def parse_upload(file_bytes: bytes, filename: str) -> ParseResult:
    """
    Extract and clean text from an uploaded resume file.

    Args:
        file_bytes: Raw bytes from a Streamlit file_uploader.
        filename:   Original filename (used to detect extension).

    Returns:
        ParseResult — always returns (never raises). Check .success.
    """
    ext = Path(filename).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        logger.warning("Unsupported extension: %s", ext)
        return ParseResult(
            is_invalid_type=True,
            error_message=f"Unsupported file type: '{ext}'",
        )

    logger.info("Parsing %s (%d bytes)", filename, len(file_bytes))

    try:
        if ext == ".pdf":
            return _parse_pdf(file_bytes)
        else:
            return _parse_docx(file_bytes)
    except Exception as exc:  # pylint: disable=broad-except
        logger.error("Unexpected error parsing %s: %s", filename, exc, exc_info=True)
        return ParseResult(
            error_message=f"Unexpected error: {exc}",
            success=False,
        )


# ── PDF extraction ─────────────────────────────────────────────────────────────

def extract_pdf_text(file_bytes: bytes) -> str:
    """
    Extract raw text from a PDF using PyMuPDF (fitz).

    Args:
        file_bytes: Raw PDF bytes.

    Returns:
        Concatenated page text, joined by newlines.

    Raises:
        ImportError: If PyMuPDF is not installed.
        fitz.FileDataError: If the file is corrupt or not a valid PDF.
    """
    import fitz  # PyMuPDF

    document = fitz.open(stream=file_bytes, filetype="pdf")
    pages = [page.get_text() for page in document]
    return "\n".join(pages).strip()


def _parse_pdf(file_bytes: bytes) -> ParseResult:
    """Internal: parse PDF with full error and quality checks."""
    try:
        import fitz
    except ImportError as exc:
        raise ImportError("PyMuPDF is required. Run: pip install pymupdf") from exc

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        return ParseResult(error_message=f"Could not open PDF: {exc}")

    # ── Encryption check ──────────────────────────────────────────────────────
    if doc.is_encrypted:
        logger.warning("PDF is password-protected")
        return ParseResult(is_encrypted=True)

    page_count = doc.page_count
    pages_text = [page.get_text() for page in doc]
    raw = "\n".join(pages_text).strip()

    # ── Scanned image check ───────────────────────────────────────────────────
    if page_count > 0:
        avg_chars = len(raw) / page_count
        if avg_chars < SCANNED_THRESHOLD:
            logger.warning(
                "Likely scanned PDF: %.0f avg chars/page across %d pages",
                avg_chars, page_count,
            )
            return ParseResult(
                is_scanned=True,
                page_count=page_count,
                char_count=len(raw),
            )

    cleaned = clean_text(raw)

    # ── Minimum content check ─────────────────────────────────────────────────
    if len(cleaned) < MIN_CHARS:
        return ParseResult(
            is_too_short=True,
            char_count=len(cleaned),
            page_count=page_count,
        )

    logger.info(
        "PDF parsed: %d chars, %d pages", len(cleaned), page_count
    )
    return ParseResult(
        text=cleaned,
        char_count=len(cleaned),
        page_count=page_count,
        success=True,
    )


# ── DOCX extraction ────────────────────────────────────────────────────────────

def extract_docx_text(file_bytes: bytes) -> str:
    """
    Extract raw text from a DOCX file using python-docx.

    Args:
        file_bytes: Raw DOCX bytes.

    Returns:
        Paragraph text joined by newlines (empty paragraphs skipped).

    Raises:
        ImportError: If python-docx is not installed.
    """
    from docx import Document

    document = Document(BytesIO(file_bytes))
    lines: list[str] = []

    # Body paragraphs.
    lines.extend(
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    )

    # Many resume templates place most content inside borderless tables.
    seen_cells: set[int] = set()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_id = id(cell._tc)  # avoid duplicated merged cells
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                lines.extend(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )

    # Preserve useful header/footer text such as a name or portfolio URL.
    seen_parts: set[int] = set()
    for section in document.sections:
        for part in (section.header, section.footer):
            part_id = id(part._element)
            if part_id in seen_parts:
                continue
            seen_parts.add(part_id)
            lines.extend(
                paragraph.text.strip()
                for paragraph in part.paragraphs
                if paragraph.text.strip()
            )

    return "\n".join(lines)


def _parse_docx(file_bytes: bytes) -> ParseResult:
    """Internal: parse DOCX with full error and quality checks."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx is required. Run: pip install python-docx") from exc

    try:
        raw = extract_docx_text(file_bytes)
    except Exception as exc:
        return ParseResult(error_message=f"Could not open DOCX: {exc}")

    cleaned = clean_text(raw)

    if len(cleaned) < MIN_CHARS:
        return ParseResult(is_too_short=True, char_count=len(cleaned))

    logger.info("DOCX parsed: %d chars", len(cleaned))
    return ParseResult(
        text=cleaned,
        char_count=len(cleaned),
        success=True,
    )


# ── Text cleaning ──────────────────────────────────────────────────────────────

def clean_text(raw: str) -> str:
    """
    Normalize and clean raw extracted resume text.

    Applies in order:
        1. Normalize unicode (NFKC) — fixes ligatures, special dashes, etc.
        2. Remove non-printable / control characters (keep newlines + tabs).
        3. Strip common page-number patterns (standalone digits on their own line).
        4. Collapse multiple blank lines into a single blank line.
        5. Collapse inline repeated spaces to a single space.
        6. Strip leading/trailing whitespace from each line.
        7. Final strip.

    Intentionally NOT removed:
        - Bullet point characters (•, ◦, ▪, –, *, -)
        - Section headings
        - Dates and numbers within sentences
    """
    if not raw:
        return ""

    # 1. Unicode normalization
    text = unicodedata.normalize("NFKC", raw)

    # 2. Remove control characters except \n and \t
    text = re.sub(r"[^\S\n\t]+", " ", text)          # runs of non-newline whitespace → single space
    text = re.sub(r"[^\x09\x0A\x20-\x7E\u00A0-\uFFFF]", "", text)  # drop stray control chars

    # 3. Remove standalone page numbers:
    #    Lines that are ONLY digits (optionally preceded by "Page" or "p.")
    text = re.sub(
        r"(?im)^\s*(page\s*)?\d{1,4}\s*$",
        "",
        text,
    )

    # 4. Collapse 3+ consecutive blank lines → 2 (preserve section spacing)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # 5. Strip trailing spaces on each line
    text = "\n".join(line.rstrip() for line in text.splitlines())

    # 6. Remove lines that are only whitespace or punctuation noise
    #    (e.g. a line with just "- - - - -" or "· · ·")
    text = re.sub(r"(?m)^[\s\-·•=_]{0,3}$", "", text)

    # 7. Collapse again after noise removal
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


# ── Backwards-compatible ResumeParser class ────────────────────────────────────

class ResumeParser:
    """
    Thin wrapper kept for backwards compatibility with existing tests.
    New code should use the module-level functions directly.
    """

    SUPPORTED_EXTENSIONS = SUPPORTED_EXTENSIONS

    def parse(self, file_path) -> str:
        """Parse a resume file by path and return clean text."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {path}")
        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type '{suffix}'.")
        result = parse_upload(path.read_bytes(), path.name)
        if not result.success:
            raise ValueError(result.user_facing_error)
        return result.text

    def parse_bytes(self, data: bytes, extension: str) -> str:
        """Parse resume from raw bytes."""
        suffix = extension.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type '{suffix}'.")
        result = parse_upload(data, f"resume{suffix}")
        if not result.success:
            raise ValueError(result.user_facing_error)
        return result.text
